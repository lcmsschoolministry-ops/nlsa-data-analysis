import os
import io
import re
import json
import asyncio
import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import Pt
from datetime import datetime
from typing import List, Dict, Any
from openai import OpenAI, AsyncOpenAI

# =========================
# Config / Defaults
# =========================
DEFAULT_MODEL_EXTRACT = "gpt-4o-mini"
DEFAULT_MODEL_SYNTH = "gpt-4o-mini"  # faster default; you can switch to gpt-4o in sidebar
MAX_SYNTH_TOKENS = 2200  # trimmed for speed

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
aclient = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

NLSA_DOMAINS = [
    "School Purpose & Mission Alignment",
    "Relationships",
    "Leadership",
    "Professional Personnel",
    "Teaching & Learning",
    "Student Services",
    "Faith Integration",
    "Operational Vitality",
]

# =========================
# PDF → Text
# =========================
def extract_pdf_text(pdf_bytes: bytes) -> List[Dict[str, Any]]:
    pages = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            txt = re.sub(r"[ \t]+", " ", txt).strip()
            pages.append({"page_number": i, "text": txt})
    return pages

def chunk_pages(pages: List[Dict[str, Any]], max_chars: int = 6000, overlap_chars: int = 250) -> List[Dict[str, Any]]:
    chunks = []
    buffer = []
    char_count = 0
    start_page = None

    def flush(end_page):
        nonlocal buffer, char_count, start_page
        if buffer:
            text = "\n".join(buffer)
            # overlap tail
            tail = text[-overlap_chars:] if overlap_chars > 0 else ""
            chunks.append({"start_page": start_page, "end_page": end_page, "text": text})
            buffer = [tail] if tail else []
            char_count = len(tail)
            start_page = None

    for p in pages:
        t = p["text"]
        if not t:
            continue
        if start_page is None:
            start_page = p["page_number"]
        if char_count + len(t) + 10 > max_chars:
            flush(p["page_number"] - 1)
            start_page = p["page_number"]
        buffer.append(f"[p.{p['page_number']}] {t}")
        char_count += len(t) + 1

    if pages:
        flush(pages[-1]["page_number"])
    return chunks

# =========================
# Prompts
# =========================
SYSTEM_EXTRACT = """You are an NLSA accreditation analyst. Extract concise, decision-useful evidence by NLSA Protocol G domains.
Output MUST be valid JSON with this schema:
{
  "domains": {
    "<domain>": {
      "evidence": ["..."],
      "strengths": ["..."],
      "growth_areas": ["..."],
      "risks_or_noncompliance": ["..."]
    }
  }
}
Limit to <= 5 bullets per list for this chunk to reduce duplication.
Include any page refs like [p.12] when present.
Domains:
- School Purpose & Mission Alignment
- Relationships
- Leadership
- Professional Personnel
- Teaching & Learning
- Student Services
- Faith Integration
- Operational Vitality
"""

USER_EXTRACT_TEMPLATE = """From the following EBA/NLSA PDF text chunk (with page refs), extract evidence aligned to NLSA Protocol G.

TEXT CHUNK:
---
{chunk_text}
---
Return ONLY JSON as specified."""

SYSTEM_SYNTH = """You are an expert NLSA (National Lutheran Schools Accreditation) reviewer using Protocol G with integrated EBA terminology.
Write a 3–4 page, administrator-ready report in professional, plain language. Use evidence-driven claims, cite page refs when present (e.g., [p.7]).
Structure:
1) Cover Block (School Name if given, Date, Report Title)
2) Executive Summary (bulleted key findings and overall judgment)
3) Domain-by-Domain Analysis (for each Protocol G domain: strengths, growth areas, notable evidence)
4) Compliance & Risk Flags (explicit mentions of potential noncompliance and missing evidence)
5) Priority Action Plan (6–12 and 12–24 month items; SMART-ish; tie to evidence)
6) Scorecard (0–100 for each domain + Composite SEI; weighting rationale)
Tone: supportive, candid, and aligned to NLSA vocabulary (e.g., “powerful practices”, “evidence”, “action plan”, “mission alignment”).
If data is missing or unclear, note it transparently and propose what evidence would satisfy Protocol G.
Target ~900–1400 words."""

USER_SYNTH_TEMPLATE = """Synthesize the FINAL REPORT from the following aggregated notes extracted across all chunks.

AGGREGATED EVIDENCE (JSON):
---
{aggregated_json}
---

Guidance:
- De-duplicate repetitive bullets.
- Prefer concrete evidence (with page refs) over generic claims.
- Assign domain scores (0–100) based on evidence: 0=absent/critical, 50=developing, 85=effective, 95+=exemplary/powerful practice.
- Compute Composite School Effectiveness Index (SEI) = average of domain scores.
- Where quantified data (enrollment, test scores, PD hours) appear, reference them.
- Faith Integration is essential; evaluate on evidence of worship, catechesis, and integration across curriculum.
Return the report as clean Markdown with headings.
"""

# =========================
# Helpers: Aggregation
# =========================
def empty_domain_dict():
    return {
        "evidence": [],
        "strengths": [],
        "growth_areas": [],
        "risks_or_noncompliance": []
    }

def merge_domain_dict(a: Dict[str, List[str]], b: Dict[str, List[str]]) -> Dict[str, List[str]]:
    out = {
        k: list(dict.fromkeys((a.get(k, []) + b.get(k, []))))  # preserve order, dedupe
        for k in ["evidence", "strengths", "growth_areas", "risks_or_noncompliance"]
    }
    return out

def aggregate_chunk_outputs(chunk_jsons: List[Dict[str, Any]], per_list_cap: int = 8) -> Dict[str, Dict[str, List[str]]]:
    agg = {d: empty_domain_dict() for d in NLSA_DOMAINS}
    for j in chunk_jsons:
        if not j or "domains" not in j:
            continue
        for domain, content in j["domains"].items():
            if domain not in agg:
                continue
            agg[domain] = merge_domain_dict(agg[domain], {
                "evidence": content.get("evidence", []),
                "strengths": content.get("strengths", []),
                "growth_areas": content.get("growth_areas", []),
                "risks_or_noncompliance": content.get("risks_or_noncompliance", []),
            })
    # Trim hard to keep synthesis fast
    for d in agg:
        for k in agg[d]:
            agg[d][k] = agg[d][k][:per_list_cap]
    return agg

# =========================
# Markdown → DOCX (minimal)
# =========================
def markdown_to_docx(md_text: str) -> Document:
    doc = Document()
    styles = doc.styles['Normal']
    styles.font.name = 'Calibri'
    styles.font.size = Pt(11)

    for line in md_text.splitlines():
        if re.match(r"^# ", line):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^## ", line):
            doc.add_heading(line[3:].strip(), level=2)
        elif re.match(r"^### ", line):
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^[-*] ", line):
            p = doc.add_paragraph(line[2:].strip())
            p.style = doc.styles['List Bullet']
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(line.strip())
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line.strip())
    return doc

# =========================
# OpenAI Calls (Async for extraction, Sync for synthesis)
# =========================
async def extract_one_async(ch_text: str, model: str, max_tokens: int = 1000) -> Dict[str, Any]:
    try:
        resp = await aclient.chat.completions.create(
            model=model,
            temperature=0.1,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": SYSTEM_EXTRACT},
                {"role": "user", "content": USER_EXTRACT_TEMPLATE.format(chunk_text=ch_text)},
            ],
            max_tokens=max_tokens,
        )
        content = resp.choices[0].message.content
        return json.loads(content)
    except Exception:
        # Fallback small text, then wrap in schema
        resp = await aclient.chat.completions.create(
            model=model,
            temperature=0.2,
            messages=[
                {"role": "system", "content": SYSTEM_EXTRACT.replace("Output MUST be valid JSON.", "Output concise bullets.")},
                {"role": "user", "content": USER_EXTRACT_TEMPLATE.format(chunk_text=ch_text)},
            ],
            max_tokens=600,
        )
        txt = resp.choices[0].message.content[:600]
        return {"domains": {d: {"evidence":[txt], "strengths":[], "growth_areas":[], "risks_or_noncompliance":[]}
                            for d in NLSA_DOMAINS}}

async def run_extractions_async(chunks: List[Dict[str, Any]], model: str, max_tokens: int, concurrency: int, progress_cb=None):
    sem = asyncio.Semaphore(concurrency)
    results = [None] * len(chunks)

    async def worker(idx: int, ch: Dict[str, Any]):
        async with sem:
            results[idx] = await extract_one_async(
                ch["text"][:6500],  # safety guard
                model=model,
                max_tokens=max_tokens
            )
            if progress_cb:
                progress_cb()

    tasks = [worker(i, ch) for i, ch in enumerate(chunks)]
    await asyncio.gather(*tasks)
    return results

def synthesize_report_sync(aggregated_json: Dict[str, Any], model: str, max_tokens: int = MAX_SYNTH_TOKENS) -> str:
    user_synth = USER_SYNTH_TEMPLATE.format(aggregated_json=json.dumps(aggregated_json, indent=2))
    resp = client.chat.completions.create(
        model=model,
        temperature=0.2,
        messages=[
            {"role": "system", "content": SYSTEM_SYNTH},
            {"role": "user", "content": user_synth},
        ],
        max_tokens=max_tokens,
    )
    return resp.choices[0].message.content

# =========================
# Chunk Prioritization (FAST MODE)
# =========================
KEYWORDS = [
    "mission", "purpose", "relationship", "govern", "board", "leadership",
    "professional", "personnel", "teaching", "learning", "curriculum",
    "assessment", "student services", "chapel", "faith", "religion",
    "budget", "finance", "enrollment", "facility", "safety", "strategic plan"
]

def score_chunk(ch: Dict[str, Any]) -> int:
    text = ch["text"].lower()
    return sum(1 for kw in KEYWORDS if kw in text)

# =========================
# Streamlit UI
# =========================
st.set_page_config(page_title="NLSA School Effectiveness (Protocol G)", page_icon="📘", layout="wide")
st.title("📘 NLSA School Effectiveness Report (Protocol G)")
st.caption("Upload an EBA/NLSA Self-Study PDF. This generates a 3–4 page report aligned to Protocol G with evidence and action plan.")

with st.sidebar:
    st.header("Settings")
    fast_mode = st.toggle("⚡ Fast Mode (limit & prioritize chunks)", value=True)
    model_extract = st.selectbox("Extraction model", [DEFAULT_MODEL_EXTRACT, "gpt-4o"], index=0)
    model_synth = st.selectbox("Synthesis model", [DEFAULT_MODEL_SYNTH, "gpt-4o"], index=0)
    if fast_mode and model_synth == "gpt-4o":
        st.info("Fast Mode prefers speed — consider keeping synthesis on gpt-4o-mini.")
    max_chars_per_chunk = st.slider("Max characters per chunk", 3000, 12000, 6000, step=500)
    overlap_chars = st.slider("Overlap characters", 0, 1000, 250, step=50)
    per_list_cap = st.slider("Max bullets per list (aggregation)", 4, 12, 8, step=1)
    max_chunks_fast = st.slider("Max chunks (Fast Mode)", 6, 20, 12, step=1)
    concurrency = st.slider("Parallel extractions (concurrency)", 2, 10, 6, step=1)
    st.markdown("---")
    st.caption("Set OPENAI_API_KEY in your environment before running.")

uploaded = st.file_uploader("Upload EBA/NLSA PDF", type=["pdf"])

if uploaded is not None:
    pdf_bytes = uploaded.read()
    with st.spinner("Extracting text from PDF..."):
        pages = extract_pdf_text(pdf_bytes)
        if not any(p["text"] for p in pages):
            st.error("No extractable text found. If the PDF is scanned images, add OCR before analysis.")
            st.stop()
    st.success(f"Parsed {len(pages)} pages.")
    st.write("Sample from first page:")
    st.code((pages[0]["text"] or "")[:800] + ("..." if len(pages[0]["text"]) > 800 else ""))

    chunks = chunk_pages(pages, max_chars=max_chars_per_chunk, overlap_chars=overlap_chars)
    st.info(f"Created {len(chunks)} chunks for analysis.")

    if fast_mode and len(chunks) > max_chunks_fast:
        chunks_scored = sorted(chunks, key=score_chunk, reverse=True)
        chunks = chunks_scored[:max_chunks_fast]
        st.warning(f"Fast Mode enabled: prioritizing top {len(chunks)} chunks by relevance keywords.")

    if st.button("🚀 Generate 3–4 Page Report"):
        # Async parallel extraction
        total = len(chunks)
        st.write(f"Analyzing {total} chunk(s) with concurrency={concurrency} …")
        progress = st.progress(0)
        counter = {"done": 0}

        def tick():
            counter["done"] += 1
            progress.progress(counter["done"] / total)

        try:
            chunk_outputs = asyncio.run(
                run_extractions_async(
                    chunks=chunks,
                    model=model_extract,
                    max_tokens=1000,  # reduced for speed
                    concurrency=concurrency,
                    progress_cb=tick
                )
            )
        except RuntimeError:
            # Fallback if an event loop is already running (rare)
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            chunk_outputs = loop.run_until_complete(
                run_extractions_async(
                    chunks=chunks,
                    model=model_extract,
                    max_tokens=1000,
                    concurrency=concurrency,
                    progress_cb=tick
                )
            )
            loop.close()

        st.success("Extraction complete. Aggregating evidence…")
        aggregated = aggregate_chunk_outputs(chunk_outputs, per_list_cap=per_list_cap)

        # Synthesis
        st.write("Synthesizing the final report…")
        report_md = synthesize_report_sync(aggregated, model=model_synth, max_tokens=MAX_SYNTH_TOKENS)

        st.subheader("📄 Report Preview")
        st.markdown(report_md)

        # DOCX export
        doc = markdown_to_docx(report_md)
        section = doc.sections[-1]
        footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer.text = f"NLSA Protocol G Report • Generated {datetime.now().strftime('%Y-%m-%d')}"
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="⬇️ Download .docx",
                data=buf,
                file_name=f"NLSA_ProtocolG_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col2:
            st.download_button(
                label="⬇️ Download Aggregated Evidence (JSON)",
                data=json.dumps(aggregated, indent=2).encode("utf-8"),
                file_name=f"NLSA_Aggregated_Evidence_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json"
            )
else:
    st.info("Upload your EBA/NLSA Self-Study PDF to begin.")
